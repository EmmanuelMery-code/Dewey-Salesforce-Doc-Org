"""Advice-section aggregation and rendering for the Word summary report.

Groups analyzer findings by rule into ``_AdviceItem`` entries and renders
them into the summary document. Extracted from ``word_writer.py`` to keep
that module focused on cover pages, tables and the data dictionary chapters.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass
from typing import TYPE_CHECKING

from docx import Document

from src.analyzer.models import SEVERITY_ORDER, Finding
from src.reporting.word_writer_labels import ADVICE_TARGET_LIMIT as _ADVICE_TARGET_LIMIT

if TYPE_CHECKING:
    from src.analyzer.engine import AnalyzerReport


@dataclass(slots=True)
class _AdviceItem:
    """Aggregated advice for a single rule across all findings.

    Sorting key: (severity rank, -occurrences) so the most severe and most
    frequent issues bubble up first.
    """

    rule_id: str
    title: str
    severity: str
    description: str
    rationale: str
    remediation: str
    targets: list[str]
    occurrences: int


class _WordAdviceMixin:
    """Adds advice-section aggregation/rendering to ``WordReportWriter``."""

    def _build_advice_items(
        self, analyzer_report: "AnalyzerReport | None"
    ) -> list[_AdviceItem]:
        if analyzer_report is None:
            return []

        findings_by_rule: dict[str, list[Finding]] = {}
        for finding in analyzer_report.all_findings():
            findings_by_rule.setdefault(finding.rule.id, []).append(finding)

        items: list[_AdviceItem] = []
        for rule_id, findings in findings_by_rule.items():
            rule = findings[0].rule
            target_counter: Counter[str] = Counter()
            for finding in findings:
                key = f"{finding.target_kind}: {finding.target_name}"
                target_counter[key] += 1
            ordered_targets = [
                target for target, _ in target_counter.most_common(_ADVICE_TARGET_LIMIT)
            ]
            items.append(
                _AdviceItem(
                    rule_id=rule_id,
                    title=rule.title or rule_id,
                    severity=rule.severity,
                    description=rule.description,
                    rationale=rule.rationale,
                    remediation=rule.remediation,
                    targets=ordered_targets,
                    occurrences=len(findings),
                )
            )

        items.sort(
            key=lambda item: (
                SEVERITY_ORDER.get(item.severity, 99),
                -item.occurrences,
                item.rule_id,
            )
        )
        return items

    def _add_advice_section(
        self, document: Document, advice: _AdviceItem, index: int
    ) -> None:
        document.add_heading(
            self._t("advice_action", index=index, title=advice.title),
            level=2,
        )

        meta_paragraph = document.add_paragraph()
        meta_paragraph.add_run(
            f"{self._t('advice_severity')}: "
        ).bold = True
        meta_paragraph.add_run(self._severity_label(advice.severity))
        meta_paragraph.add_run("    ")
        meta_paragraph.add_run(
            f"{self._t('advice_occurrences')}: "
        ).bold = True
        meta_paragraph.add_run(str(advice.occurrences))

        if advice.description:
            self._add_labelled_paragraph(
                document, self._t("advice_description"), advice.description
            )
        if advice.rationale:
            self._add_labelled_paragraph(
                document, self._t("advice_rationale"), advice.rationale
            )
        if advice.remediation:
            self._add_labelled_paragraph(
                document, self._t("advice_remediation"), advice.remediation
            )

        if advice.targets:
            document.add_paragraph(self._t("advice_examples"), style="Heading 3")
            for target in advice.targets:
                document.add_paragraph(target, style="List Bullet")
            remaining = advice.occurrences - len(advice.targets)
            if remaining > 0:
                document.add_paragraph(
                    self._t("advice_examples_more", count=remaining)
                ).runs[0].italic = True

    def _severity_label(self, severity: str) -> str:
        mapping = {
            "Critical": self._t("severity_critical"),
            "Major": self._t("severity_major"),
            "Minor": self._t("severity_minor"),
            "Info": self._t("severity_info"),
        }
        return mapping.get(severity, severity)
