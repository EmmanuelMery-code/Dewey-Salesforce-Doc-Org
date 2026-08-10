from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Callable

LogCallback = Callable[[str], None]

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from src.core.models import MetadataSnapshot, ObjectInfo
from src.reporting.word_writer_advice import _WordAdviceMixin
from src.reporting.word_writer_labels import LABELS as _LABELS

if TYPE_CHECKING:
    from src.analyzer.engine import AnalyzerReport


class WordReportWriter(_WordAdviceMixin):
    """Generates the Word counterparts of the documentation.

    The writer is intentionally decoupled from the rest of the reporting
    pipeline: it only consumes plain dataclasses (`MetadataSnapshot`,
    `AnalyzerReport`) and writes ``.docx`` files into the directory that
    callers specify. Advice-section aggregation/rendering lives in
    :class:`~src.reporting.word_writer_advice._WordAdviceMixin`.
    """

    def __init__(
        self, language: str = "fr", log_callback: LogCallback | None = None
    ) -> None:
        self.language = language if language in _LABELS else "fr"
        self.log: LogCallback = log_callback or (lambda message: None)

    # ------------------------------------------------------------------ public API

    def write_data_dictionary_document(
        self,
        snapshot: MetadataSnapshot,
        output_path: str | Path,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Excluded objects are already filtered upstream in the parser; we
        # additionally drop objects that have no field at all so the
        # generated document mirrors the Excel dictionary exactly.
        documented_objects = [obj for obj in snapshot.objects if obj.fields]

        document = Document()
        self._configure_default_style(document)
        self._enable_field_auto_update(document)

        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
        self._add_cover_page(
            document,
            title=self._t("data_dictionary_doc_title"),
            subtitle=self._t("data_dictionary_subtitle", date=generated_at),
        )

        self._add_table_of_contents_page(document)

        if not documented_objects:
            document.add_paragraph(self._t("no_objects"))
        else:
            total = len(documented_objects)
            for index, obj in enumerate(documented_objects):
                if index > 0:
                    document.add_page_break()
                
                if index % 10 == 0:
                    self.log(f"Generation Word : objet {index + 1}/{total} ({obj.api_name})")
                
                self._add_object_chapter(document, obj)

        document.save(output_path)
        self.log(
            f"Data Dictionary Word genere ({len(documented_objects)} objet(s)) : "
            f"{output_path}"
        )
        return output_path

    def write_summary_document(
        self,
        snapshot: MetadataSnapshot,
        analyzer_report: "AnalyzerReport | None",
        output_path: str | Path,
    ) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        self._configure_default_style(document)

        generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")
        self._add_cover_page(
            document,
            title=self._t("summary_doc_title"),
            subtitle=self._t("summary_subtitle", date=generated_at),
        )

        document.add_heading(self._t("section_overview"), level=1)
        document.add_paragraph(self._t("overview_intro"))

        document.add_heading(self._t("section_metrics"), level=1)
        document.add_paragraph(self._t("overview_metrics_intro"))
        self._add_metrics_table(document, snapshot, analyzer_report)

        document.add_heading(self._t("section_advice"), level=1)
        advice_items = self._build_advice_items(analyzer_report)
        if not advice_items:
            document.add_paragraph(self._t("advice_no_findings"))
        else:
            document.add_paragraph(self._t("advice_intro"))
            for index, advice in enumerate(advice_items, start=1):
                self._add_advice_section(document, advice, index)

        document.save(output_path)
        self.log(f"Resume Word genere : {output_path}")
        return output_path

    # ------------------------------------------------------------------ helpers

    def _t(self, key: str, **fmt: object) -> str:
        labels = _LABELS.get(self.language) or _LABELS["fr"]
        template = labels.get(key, key)
        if not fmt:
            return template
        try:
            return template.format(**fmt)
        except (KeyError, IndexError):
            return template

    @staticmethod
    def _configure_default_style(document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

    @staticmethod
    def _enable_field_auto_update(document: Document) -> None:
        # Tells Word to prompt the user to update fields (i.e. the TOC) when
        # the document is opened. Without this the TOC stays empty until the
        # user manually presses F9.
        settings = document.settings.element
        existing = settings.find(qn("w:updateFields"))
        if existing is None:
            update_fields = OxmlElement("w:updateFields")
            update_fields.set(qn("w:val"), "true")
            settings.append(update_fields)

    def _add_cover_page(self, document: Document, *, title: str, subtitle: str) -> None:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = 1  # center
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(28)

        subtitle_paragraph = document.add_paragraph()
        subtitle_paragraph.alignment = 1
        subtitle_run = subtitle_paragraph.add_run(subtitle)
        subtitle_run.font.size = Pt(16)

        # Force the cover to occupy a full page so the table of contents
        # naturally lands on page 2.
        document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    def _add_table_of_contents_page(self, document: Document) -> None:
        document.add_heading(self._t("table_of_contents"), level=1)
        hint = document.add_paragraph(self._t("table_of_contents_hint"))
        for run in hint.runs:
            run.italic = True

        toc_paragraph = document.add_paragraph()
        run = toc_paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        # h: hyperlinks, z: hide tab leader on web, u: use heading styles.
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_separate)
        run._r.append(fld_end)

        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def _add_object_chapter(self, document: Document, obj: ObjectInfo) -> None:
        if obj.label:
            heading = self._t(
                "object_chapter_title",
                label=obj.label,
                api_name=obj.api_name,
            )
        else:
            heading = self._t("object_chapter_title_simple", api_name=obj.api_name)
        document.add_heading(heading, level=1)

        document.add_heading(self._t("section_information"), level=2)
        self._add_information_table(document, obj)

        document.add_heading(self._t("section_fields"), level=2)
        self._add_fields_table(document, obj)

    def _add_information_table(self, document: Document, obj: ObjectInfo) -> None:
        rows: list[tuple[str, str]] = [
            (self._t("info_api_name"), obj.api_name or self._t("value_unspecified")),
            (self._t("info_label"), obj.label or self._t("value_unspecified")),
            (
                self._t("info_plural_label"),
                obj.plural_label or self._t("value_unspecified"),
            ),
            (
                self._t("info_custom"),
                self._t("yes") if obj.custom else self._t("no"),
            ),
            (
                self._t("info_sharing_model"),
                obj.sharing_model or self._t("value_unspecified"),
            ),
            (
                self._t("info_deployment_status"),
                obj.deployment_status or self._t("value_unspecified"),
            ),
            (
                self._t("info_visibility"),
                obj.visibility or self._t("value_unspecified"),
            ),
            (self._t("info_field_count"), str(len(obj.fields))),
            (
                self._t("info_custom_field_count"),
                str(sum(1 for field in obj.fields if field.custom)),
            ),
            (self._t("info_record_types"), str(len(obj.record_types))),
            (self._t("info_validation_rules"), str(len(obj.validation_rules))),
            (self._t("info_relationships"), str(len(obj.relationships))),
            (
                self._t("info_description"),
                (obj.description or "").strip() or self._t("value_unspecified"),
            ),
        ]
        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for index, (label, value) in enumerate(rows):
            label_cell = table.cell(index, 0)
            value_cell = table.cell(index, 1)
            self._set_cell_text(label_cell, label, bold=True)
            self._set_cell_text(value_cell, value)
        self._set_table_column_widths(table, [Cm(5.5), Cm(11.0)])

    def _add_fields_table(self, document: Document, obj: ObjectInfo) -> None:
        headers = [
            self._t("field_column_label"),
            self._t("field_column_api_name"),
            self._t("field_column_type"),
            self._t("field_column_description"),
        ]
        table = document.add_table(rows=1 + len(obj.fields), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for col_index, header in enumerate(headers):
            self._set_cell_text(table.cell(0, col_index), header, bold=True)

        sorted_fields = sorted(
            obj.fields,
            key=lambda f: ((f.label or f.api_name or "").lower()),
        )
        for row_index, field_info in enumerate(sorted_fields, start=1):
            description = (field_info.description or "").strip()
            if not description:
                description = self._t("field_no_description")
            self._set_cell_text(
                table.cell(row_index, 0),
                field_info.label or field_info.api_name,
            )
            self._set_cell_text(table.cell(row_index, 1), field_info.api_name)
            self._set_cell_text(
                table.cell(row_index, 2),
                field_info.data_type or self._t("value_unspecified"),
            )
            self._set_cell_text(table.cell(row_index, 3), description)

        self._set_table_column_widths(
            table,
            [Cm(4.5), Cm(4.5), Cm(2.5), Cm(5.0)],
        )

    def _add_metrics_table(
        self,
        document: Document,
        snapshot: MetadataSnapshot,
        analyzer_report: "AnalyzerReport | None",
    ) -> None:
        metrics = snapshot.metrics
        rows: list[tuple[str, str]] = [
            (self._t("metric_objects"), str(len(snapshot.objects))),
            (self._t("metric_custom_objects"), str(metrics.custom_objects)),
            (self._t("metric_custom_fields"), str(metrics.custom_fields)),
            (self._t("metric_record_types"), str(metrics.record_types)),
            (self._t("metric_validation_rules"), str(metrics.validation_rules)),
            (self._t("metric_layouts"), str(metrics.layouts)),
            (self._t("metric_custom_tabs"), str(metrics.custom_tabs)),
            (self._t("metric_custom_apps"), str(metrics.custom_apps)),
            (self._t("metric_flows"), str(metrics.flows)),
            (self._t("metric_apex_classes"), str(metrics.apex_classes)),
            (self._t("metric_apex_triggers"), str(metrics.apex_triggers)),
            (self._t("metric_lwc"), str(metrics.lwc_count)),
            (self._t("metric_flexipages"), str(metrics.flexipage_count)),
            (self._t("metric_omni_scripts"), str(metrics.omni_scripts)),
            (
                self._t("metric_omni_integration_procedures"),
                str(metrics.omni_integration_procedures),
            ),
            (self._t("metric_omni_ui_cards"), str(metrics.omni_ui_cards)),
            (
                self._t("metric_omni_data_transforms"),
                str(metrics.omni_data_transforms),
            ),
            (self._t("metric_score"), str(metrics.score)),
            (self._t("metric_level"), metrics.level),
            (self._t("metric_adopt_adapt_score"), str(metrics.adopt_adapt_score)),
            (self._t("metric_adopt_adapt_level"), metrics.adopt_adapt_level),
        ]
        if analyzer_report is not None:
            severity_counts = analyzer_report.severity_counts()
            total = sum(severity_counts.values())
            rows.extend(
                [
                    (self._t("metric_findings_total"), str(total)),
                    (
                        self._t("metric_findings_critical"),
                        str(severity_counts.get("Critical", 0)),
                    ),
                    (
                        self._t("metric_findings_major"),
                        str(severity_counts.get("Major", 0)),
                    ),
                    (
                        self._t("metric_findings_minor"),
                        str(severity_counts.get("Minor", 0)),
                    ),
                    (
                        self._t("metric_findings_info"),
                        str(severity_counts.get("Info", 0)),
                    ),
                ]
            )

        table = document.add_table(rows=len(rows), cols=2)
        table.style = "Light Grid Accent 1"
        for index, (label, value) in enumerate(rows):
            self._set_cell_text(table.cell(index, 0), label, bold=True)
            self._set_cell_text(table.cell(index, 1), value)
        self._set_table_column_widths(table, [Cm(7.0), Cm(8.0)])

    @staticmethod
    def _add_labelled_paragraph(document: Document, label: str, body: str) -> None:
        cleaned = " ".join((body or "").split())
        if not cleaned:
            return
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(cleaned)

    @staticmethod
    def _set_cell_text(cell, value: str, *, bold: bool = False) -> None:
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(value)
        run.bold = bold

    @staticmethod
    def _set_table_column_widths(table, widths) -> None:
        for row in table.rows:
            for index, width in enumerate(widths):
                if index < len(row.cells):
                    row.cells[index].width = width
